
import sys
import os
from docx import Document

def inspect_pagination_clues(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    print(f"Inspecting Layout Clues: {docx_path}")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Check the last few elements
    print("\n[End of Document Check]")
    for i, p in enumerate(doc.paragraphs[-10:]):
        idx = len(doc.paragraphs) - 10 + i
        text = p.text.strip()
        print(f"Para {idx}: '{text}' (Length: {len(text)})")
        if not text:
             print("  -> WARNING: Empty Paragraph (Potential blank page cause)")

    # Check Table 2 (Self Intro) length
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        print(f"\n[Table 2 - Self Intro]")
        rows = len(t2.rows)
        print(f"Rows: {rows}")
        # Assuming content is in row 1, col 0 based on fill_doc.py
        if rows > 1:
            cell = t2.cell(1, 0)
            text = cell.text
            print(f"Content Length: {len(text)} chars")
            print(f"Line count (approx): {text.count('\n') + 1}")
            print(f"Preview: {text[:50]}...")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python inspect_layout.py <path_to_docx>")
        sys.exit(1)
    inspect_pagination_clues(sys.argv[1])
