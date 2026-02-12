"""v9 → v10 패치: 사진칸에 프로필 사진 삽입 (증명사진 비율 크롭)"""
import os
import sys
import tempfile
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PROJECT_ROOT = os.path.abspath(os.path.join(BASE_DIR, '..', '..', '..'))
INPUT = os.path.join(BASE_DIR, 'drafts', 'resume_v9.docx')
OUTPUT = os.path.join(BASE_DIR, 'drafts', 'resume_v10.docx')
PHOTO_PATH = os.path.join(PROJECT_ROOT, 'common', 'photos', 'profile.png')

# 사진 셀: R0C0, 수직병합 R0~R2, 폭 2000 twips ≈ 3.5cm
# 증명사진 표준 비율 3:4 사용, 셀 폭에 맞춤
PHOTO_WIDTH_CM = 3.0   # 셀 3.5cm에서 마진 빼고
PHOTO_RATIO = 3.0 / 4.0  # 가로:세로 = 3:4 증명사진


def crop_portrait(photo_path, out_path):
    """3:4 증명사진 비율로 중앙 크롭 (얼굴 상단 우선)"""
    with Image.open(photo_path) as img:
        img = img.convert('RGB')
        w, h = img.size
        src_ratio = w / h

        if src_ratio > PHOTO_RATIO:
            new_w = int(h * PHOTO_RATIO)
            left = (w - new_w) // 2
            img = img.crop((left, 0, left + new_w, h))
        elif src_ratio < PHOTO_RATIO:
            new_h = int(w / PHOTO_RATIO)
            top = max(0, int((h - new_h) * 0.15))
            if top + new_h > h:
                top = h - new_h
            img = img.crop((0, top, w, top + new_h))

        img = img.resize((480, 640), Image.LANCZOS)
        img.save(out_path, 'PNG')
    return out_path


def clear_cell(cell):
    """셀 내 기존 내용(텍스트+이미지) 모두 제거"""
    for p in cell.paragraphs:
        # 인라인 이미지 제거
        for drawing in p._element.findall(qn('w:r') + '/' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        for r in p.runs:
            r.clear()
        p.text = ''


def main():
    if not os.path.exists(INPUT):
        print(f"Error: {INPUT} not found"); sys.exit(1)
    if not os.path.exists(PHOTO_PATH):
        print(f"Error: {PHOTO_PATH} not found"); sys.exit(1)

    doc = Document(INPUT)
    cell = doc.tables[0].cell(0, 0)

    # 기존 내용 제거
    clear_cell(cell)

    # 사진 크롭
    cropped = os.path.join(tempfile.gettempdir(), 'sk_profile_cropped.png')
    crop_portrait(PHOTO_PATH, cropped)
    print(f"크롭 완료 (3:4 비율, 480x640)")

    # 삽입
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 수직 가운데 정렬
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign_elem = etree.SubElement(tcPr, qn('w:vAlign'))
        vAlign_elem.set(qn('w:val'), 'center')
    else:
        vAlign.set(qn('w:val'), 'center')

    run = paragraph.add_run()
    run.add_picture(cropped, width=Cm(PHOTO_WIDTH_CM))
    print(f"사진 삽입: {PHOTO_WIDTH_CM}cm 폭")

    doc.save(OUTPUT)
    print(f"\n저장: {OUTPUT}")


if __name__ == '__main__':
    main()
