
import sys
import os
from docx import Document
from docx.shared import Pt

def inspect_breaks(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    print(f"Inspecting breaks in: {docx_path}")
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")

    # Iterate through all block-level elements (paragraphs and tables are siblings in XML but docx API treats them separately)
    # python-docx doesn't easily show the order of tables vs paragraphs.
    # We will check paragraphs for page_break_before property.

    print("\n--- Paragraph Inspection ---")
    for i, p in enumerate(doc.paragraphs):
        info = []
        if p.paragraph_format.page_break_before:
            info.append("PAGE_BREAK_BEFORE")
        
        # Check for hard page breaks in runs
        for r in p.runs:
            if 'lastRenderedPageBreak' in r._element.xml:
                 info.append("Run:lastRenderedPageBreak")
            if 'br' in r._element.xml and 'type="page"' in r._element.xml:
                 info.append("Run:HARD_BREAK")

        text = p.text.strip()
        if info or not text:
            print(f"Para {i}: '{text[:20]}...' flags={info}")
    
    # Check Table properties (not easily accessible via pure python-docx for page breaks without OXML) but we can try.
    print("\n--- Table Inspection ---")
    pass # Tables mostly don't have direct page break properties in high level API, usually it's the paragraph inside or before.

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python inspect_breaks.py <path_to_docx>")
        sys.exit(1)
    inspect_breaks(sys.argv[1])
