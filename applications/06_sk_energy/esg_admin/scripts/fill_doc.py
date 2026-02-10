"""
Word (.docx) 이력서 양식 자동 입력 도구

사용법:
    python tools/fill_doc.py <input.docx> <output.docx> [--profile content/profile.json]

의존성: python-docx, Pillow
"""

import json
import os
import sys
import argparse
import tempfile
from datetime import date
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


def load_profile(profile_path):
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def set_cell(table, row, col, value, font_name='맑은 고딕', font_size=None, bold=False):
    """셀 텍스트 교체 (서식 보존 + 폰트 강제 적용). 0-indexed grid 좌표."""
    try:
        cell = table.cell(row, col)
        # 기존 내용 삭제
        cell.text = str(value)
        
        # 스타일 적용
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = font_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if font_size:
                    run.font.size = Pt(font_size)
                if bold:
                    run.font.bold = True
    except (IndexError, AttributeError) as e:
        print(f"Error setting cell ({row}, {col}): {e}")


def insert_photo(table, row, col, photo_path, width_cm=2.5):
    """셀에 사진 삽입."""
    try:
        cell = table.cell(row, col)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(photo_path, width=Cm(width_cm))
    except Exception as e:
        print(f'사진 삽입 실패: {e}')

def allow_row_break(table):
    """표의 모든 행이 페이지를 넘어갈 수 있도록 허용."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # cantSplit 요소가 있으면 삭제 (페이지 넘기기 허용)
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is not None:
            trPr.remove(cantSplit)


def fill_resume(doc, profile, project_root):
    """간편이력서 양식 (Table 0) 자동 입력"""
    t = doc.tables[0]
    allow_row_break(t)  # 페이지 넘김 허용 설정
    p = profile['personal_info']

    # 사진
    photo_path = os.path.join(project_root, 'common', 'photos', 'profile.png')
    if os.path.exists(photo_path):
        insert_photo(t, 0, 0, photo_path, width_cm=2.5)

    # 응시파트
    set_cell(t, 2, 18, 'ESG 사무', font_size=11, bold=True)

    # 인적사항
    set_cell(t, 4, 5, p['name_ko'], font_size=10)
    set_cell(t, 4, 15, p['birthdate'].replace('-', '.'), font_size=10)
    set_cell(t, 5, 5, p['email'], font_size=9)  # 이메일은 길 수 있으므로 9pt
    set_cell(t, 5, 15, p['phone'], font_size=10)
    set_cell(t, 6, 5, p['address'], font_size=9) # 주소도 길 수 있음

    # 학력사항
    edu = profile['education']

    # R9: 고등학교
    hs = next((e for e in edu if '고등학교' in e.get('school', '')), None)
    if hs:
        period = hs.get('period', '').split(' ~ ')
        if len(period) == 2:
            set_cell(t, 9, 1, period[0], font_size=9)
            set_cell(t, 9, 6, period[1], font_size=9)
        set_cell(t, 9, 8, hs['school'].replace('고등학교', ''), font_size=10)
        set_cell(t, 9, 19, '졸업', font_size=10)

    # R10: 대학
    undergrad = next((e for e in edu if e.get('degree') == '학사'), None)
    if undergrad:
        period = undergrad.get('period', '').split(' ~ ')
        if len(period) == 2:
            set_cell(t, 10, 1, period[0], font_size=9)
            set_cell(t, 10, 6, period[1], font_size=9)
        set_cell(t, 10, 8, '숭실', font_size=10)
        set_cell(t, 10, 13, '전기공학/경제학(복수)', font_size=9) # 전공명 길어서 9pt
        set_cell(t, 10, 19, '졸업', font_size=10)
        set_cell(t, 10, 22, undergrad.get('gpa', ''), font_size=10)

    # R11: 대학원
    grad = next((e for e in edu if e.get('degree') == '석사'), None)
    if grad:
        period = grad.get('period', '').split(' ~ ')
        if len(period) == 2:
            set_cell(t, 11, 1, period[0], font_size=9)
            set_cell(t, 11, 6, period[1], font_size=9)
        set_cell(t, 11, 8, '숭실 사회복지', font_size=10)
        set_cell(t, 11, 13, '사회적기업', font_size=10)
        set_cell(t, 11, 19, '졸업', font_size=10)
        set_cell(t, 11, 22, grad.get('gpa', ''), font_size=10)

    # 경력사항 (R14~R16) — 간결하게
    career_data = [
        ('2023.01', '2026.02', '(사)함께만드는세상', '총무·IT지원(대리)', '3,423만', '자발적 퇴사'),
        ('2020.03', '2021.11', '(주)맨지온', '임베디드개발/QA(대리)', '3,600만', '계약만료'),
        ('2016.05', '2018.02', '(주)엣지크로스', 'H/W개발·자재관리(주임)', '2,760만', '계약만료'),
    ]
    for i, (start, end, company, role, salary, reason) in enumerate(career_data):
        row = 14 + i
        set_cell(t, row, 1, start, font_size=9)
        set_cell(t, row, 6, end, font_size=9)
        set_cell(t, row, 8, company, font_size=9) # 회사명 길이 고려
        set_cell(t, row, 12, role, font_size=8)   # 직무명 길어서 8pt
        set_cell(t, row, 17, salary, font_size=9)
        set_cell(t, row, 21, reason, font_size=9)

    # PC 활용능력 (R20~R23)
    set_cell(t, 20, 4, '상', font_size=10)
    set_cell(t, 21, 4, '상', font_size=10)
    set_cell(t, 22, 4, '상', font_size=10)
    set_cell(t, 23, 4, 'SQL, AI Tools', font_size=8) # 내용이 길어서 8pt

    # 자격증
    certs = [('사회복지사 1급', '2020.02'), ('1종대형 운전면허', '2019.02')]
    for i, (name, date) in enumerate(certs):
        set_cell(t, 20 + i, 7, name, font_size=9)
        set_cell(t, 20 + i, 9, date, font_size=9)

    # 병역사항 (profile.json 'military' 필드 참조)
    mil = profile.get('military', {})
    if mil:
        # Row 20: 군별 (Service Type)
        set_cell(t, 20, 20, mil.get('service_type', ''), font_size=10)
        # Row 21: 계급 (Rank)
        set_cell(t, 21, 20, mil.get('rank', ''), font_size=10)
        # Row 22: 복무기간 (Period)
        set_cell(t, 22, 20, mil.get('period', ''), font_size=7) # 기간이 길어서 7pt로 축소
        # Row 23: 면제사유 (만기제대인 경우 공란 또는 '만기제대')
        if mil.get('discharge_type') != '만기제대':
            set_cell(t, 23, 20, mil.get('discharge_type', ''), font_size=9)
    else:
        # Fallback if no military data
        set_cell(t, 20, 20, '군필', font_size=10)


def fill_etc(doc):
    """기타사항 (Table 1)"""
    t = doc.tables[1]
    set_cell(t, 1, 0,
             '* 출퇴근 약 50분 소요 예상 * 즉시 면접( 가능 ) 및 출근( 가능 )')


def generate_seal(name, output_path, size=200):
    """빨간 원형 도장 이미지 생성."""
    img = Image.new('RGBA', (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # 빨간 원 테두리
    margin = 8
    draw.ellipse(
        [margin, margin, size - margin, size - margin],
        outline=(200, 30, 30, 255), width=6
    )

    # 이름 텍스트 (세로 배치: 임 / 종 / 권)
    try:
        font = ImageFont.truetype('C:/Windows/Fonts/malgunbd.ttf', size // 4)
    except OSError:
        font = ImageFont.load_default()

    chars = list(name)
    total_h = len(chars) * (size // 4 + 4)
    y_start = (size - total_h) // 2 + 4
    for i, ch in enumerate(chars):
        bbox = draw.textbbox((0, 0), ch, font=font)
        tw = bbox[2] - bbox[0]
        x = (size - tw) // 2
        y = y_start + i * (size // 4 + 4)
        draw.text((x, y), ch, fill=(200, 30, 30, 255), font=font)

    img.save(output_path, 'PNG')
    return output_path


def fill_date_signature(doc, name):
    """날짜 및 서명(도장) 입력. P18: '20__년__월__일 지원자:___(인)'"""
    p = doc.paragraphs[18]
    today = date.today()

    # Run0:"20 " → "2026"  Run3:"     " → " 02"  Run5:"     " → " 09"  Run9:"            " → " 임종권 "
    if len(p.runs) >= 11:
        p.runs[0].text = str(today.year)
        p.runs[1].text = ''
        p.runs[3].text = f' {today.month:02d}'
        p.runs[5].text = f' {today.day:02d}'
        p.runs[9].text = f' {name} '

        # (인) 자리에 도장 이미지 삽입
        seal_path = os.path.join(tempfile.gettempdir(), 'seal_stamp.png')
        generate_seal(name, seal_path, size=140)
        p.runs[10].text = ''
        p.runs[10].add_picture(seal_path, width=Cm(1.2))


def fill_self_intro(doc):
    """자기소개서 (Table 2) — ESG 사무직 맞춤"""
    t = doc.tables[2]
    intro = (
        "전기공학 학사와 사회복지학 석사(사회적기업 전공)를 거쳐, "
        "비영리 법인에서 3년간 총무·IT지원 및 사회공헌사업 운영 실무를 담당했습니다.\n\n"

        "■ 사회공헌 사업 운영 경험\n"
        "(사)함께만드는세상에서 KDB나눔재단의 '시니어브리지 사회공헌아카데미' "
        "사업 운영에 참여했습니다. 교육생 모집·접수, 외부 협력기관 10개+ 선정·계약·관리, "
        "교육장 환경 조성, 강사 섭외 보조 등 사업 전 주기를 지원했습니다. "
        "협력기관 관리 프로세스를 표준화하여 정산 보고서 작성 시간을 30% 단축했습니다.\n\n"

        "■ 예산 관리 및 비용 절감\n"
        "인프라지원팀 배정 예산 약 3.6억 원의 집행·통제를 담당했습니다(집행률 88.5%). "
        "복합기 렌탈, 전자계약서비스 등 재계약 협상을 통해 연간 약 500만 원 비용을 절감했습니다.\n\n"

        "■ IT 활용 및 ERP 도입\n"
        "커스텀 ERP(Tibero/Nexacro 기반) 도입 프로젝트를 실무 전담하여, "
        "요구사항 정의부터 회계·인사 프로세스 설계, 검수까지 수행했습니다. "
        "SQL 데이터 추출·분석, Excel 피벗/VBA, PPT, AI 도구 활용에 능숙합니다.\n\n"

        "■ 지원 동기\n"
        "비영리 현장에서 사회적 가치를 추구하는 일에 보람을 느꼈고, "
        "ESG·사회공헌 분야에서 실무 역량을 발전시키고 싶습니다. "
        "사업 운영, 예산 관리, 외부 커뮤니케이션 경험을 바탕으로 "
        "맡겨주신 업무를 책임감 있게 수행하겠습니다."
    )
    set_cell(t, 1, 0, intro)


def remove_hard_breaks(doc):
    """문서 내의 강제 페이지 나누기(Hard Page Break) 및 단락 전 페이지 넘김 제거"""
    # 1. 단락 속성: page_break_before 제거
    for p in doc.paragraphs:
        if p.paragraph_format.page_break_before:
            p.paragraph_format.page_break_before = False

    # 2. Run 내의 <w:br w:type="page"/> 제거
    # python-docx에서는 run.break()가 <w:br>을 생성함.
    # 이를 제거하기 위해 XML 직접 조작
    for p in doc.paragraphs:
        for run in p.runs:
            if 'br' in run._element.xml:
                brs = run._element.findall(qn('w:br'))
                for br in brs:
                    if br.get(qn('w:type')) == 'page':
                        run._element.remove(br)

def cleanup_trailing_paragraphs(doc):
    """문서 끝의 빈 단락 제거 (빈 페이지 방지)"""
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]
        if not p.text.strip() and not p.runs:
            # XML 요소 삭제
            p._element.getparent().remove(p._element)
        else:
            break

def main():
    parser = argparse.ArgumentParser(description='Word 이력서 양식 자동 입력')
    parser.add_argument('input', help='입력 .docx 파일 경로')
    parser.add_argument('output', help='출력 .docx 파일 경로')
    parser.add_argument('--profile', default='content/profile.json', help='profile.json 경로')
    args = parser.parse_args()

    input_path = os.path.abspath(args.input)
    output_path = os.path.abspath(args.output)
    profile_path = os.path.abspath(args.profile)
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    if not os.path.exists(input_path):
        print(f'Error: {input_path} not found')
        sys.exit(1)

    profile = load_profile(profile_path)
    doc = Document(input_path)

    fill_resume(doc, profile, project_root)
    fill_etc(doc)
    fill_self_intro(doc)
    fill_date_signature(doc, profile['personal_info']['name_ko'])
    
    # 마지막 정리
    remove_hard_breaks(doc)
    cleanup_trailing_paragraphs(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'저장 완료: {output_path}')


if __name__ == '__main__':
    main()
