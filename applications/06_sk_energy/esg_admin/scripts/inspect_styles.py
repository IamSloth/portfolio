
import sys
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_alignment_name(alignment_val):
    if alignment_val == WD_ALIGN_PARAGRAPH.CENTER: return "CENTER"
    if alignment_val == WD_ALIGN_PARAGRAPH.LEFT: return "LEFT"
    if alignment_val == WD_ALIGN_PARAGRAPH.RIGHT: return "RIGHT"
    if alignment_val == WD_ALIGN_PARAGRAPH.JUSTIFY: return "JUSTIFY"
    if alignment_val is None: return "NONE(Left)"
    return str(alignment_val)

def inspect_table_styles(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    print(f"Inspecting Styles: {docx_path}")
    
    if not doc.tables:
        print("No tables found.")
        return
        
    table = doc.tables[0]
    
    # Target Rows: 
    # Row 9 (High School) - Check alignment of School Name
    # Row 20-22 (Military) - Check alignment of details
    
    target_rows = [9, 10, 11, 20, 21, 22]
    
    for r_idx in target_rows:
        if r_idx >= len(table.rows):
            continue
        row = table.rows[r_idx]
        print(f"\n[Row {r_idx}]")
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip().replace('\n', ' ')
            if not text: continue
            
            # Check paragraph formatting
            p = cell.paragraphs[0]
            align = get_alignment_name(p.alignment)
            
            # Check font
            font_info = []
            for run in p.runs:
                if run.font.name:
                    font_info.append(run.font.name)
                if run.font.size:
                    font_info.append(f"{run.font.size.pt}pt")
            
            font_str = ", ".join(font_info) if font_info else "Default Font"
            
            print(f"  Col {c_idx}: '{text[:15]}...' | Align: {align} | Font: {font_str}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python inspect_styles.py <path_to_docx>")
        sys.exit(1)
    inspect_table_styles(sys.argv[1])
