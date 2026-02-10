"""v8 → v9 패치: 자기소개서 좌측정렬 + 내용 수정 2건"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT = os.path.join(os.path.dirname(__file__), '..', 'applications', 'sk_energy', 'esg_admin', 'drafts', 'resume_v8.docx')
OUTPUT = INPUT.replace('v8', 'v9')

doc = Document(os.path.abspath(INPUT))

# --- 1. 자기소개서 (Table 2) 본문 좌측 정렬 ---
t = doc.tables[2]
for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            text = p.text.strip()
            # 제목("자 기 소 개 서")만 가운데 유지, 나머지 좌측
            if text and '자 기 소 개' not in text:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 2. 텍스트 치환 (전체 문서 paragraphs + tables) ---
REPLACEMENTS = {
    "커스텀 ERP(Tibero/Nexacro 기반) 도입 프로젝트를 실무 전담하여":
        "전사 ERP 시스템 도입 프로젝트를 실무 전담하여",
    "비영리 현장에서 사회적 가치를 추구하는 일에 보람을 느꼈고, ESG·사회공헌 분야에서 실무 역량을 발전시키고 싶습니다.":
        "비영리 현장에서 사회공헌사업 운영 실무를 담당하며 이 분야에서 성장하겠다는 확신을 갖게 되었습니다. SK그룹의 사회적 가치 추구 경영철학에 공감하며, 현장 경험을 바탕으로 ESG 사회공헌 업무에 기여하고 싶습니다.",
}

def replace_in_paragraphs(paragraphs):
    count = 0
    for p in paragraphs:
        full = p.text
        modified = full
        matched = False
        for old, new in REPLACEMENTS.items():
            if old in modified:
                modified = modified.replace(old, new)
                matched = True
                count += 1
        if matched and p.runs:
            p.runs[0].text = modified
            for r in p.runs[1:]:
                r.text = ''
    return count

replaced = 0
# Table 내 paragraphs
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replaced += replace_in_paragraphs(cell.paragraphs)
# 문서 본문 paragraphs
replaced += replace_in_paragraphs(doc.paragraphs)

doc.save(os.path.abspath(OUTPUT))
print(f"저장 완료: {OUTPUT}")
print(f"텍스트 치환: {replaced}건")
