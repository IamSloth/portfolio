
import sys
import os
from docx import Document

def inspect_table_structure(docx_path):
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    print(f"Inspecting: {docx_path}")
    
    # Focus on Table 0 (Main Resume)
    if not doc.tables:
        print("No tables found.")
        return
        
    table = doc.tables[0]
    print(f"\n[Table 0] Total Rows: {len(table.rows)}")
    
    # Inspect rows 19 to 23 (Military Service Area)
    # Adjust range if needed based on previous output ensuring we cover the military section
    target_rows = range(19, 24) 
    
    for r_idx in target_rows:
        if r_idx >= len(table.rows):
            break
        row = table.rows[r_idx]
        row_content = []
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip().replace('\n', ' ')
            # Only print if text exists or it's a target column we are suspicious of
            if text:
                row_content.append(f"[{c_idx}]='{text}'")
        
        print(f"Row {r_idx}: " + " | ".join(row_content))

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python inspect_docx.py <path_to_docx>")
        sys.exit(1)
    inspect_table_structure(sys.argv[1])
