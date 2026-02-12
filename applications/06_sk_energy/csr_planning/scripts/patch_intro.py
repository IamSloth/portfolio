"""ESG 사무용 resume_v10.docx → 사회공헌 기획/운영용 resume_v1.docx
자기소개서(Table 2, Cell(1,0))만 교체"""
import os
import sys
import shutil
from docx import Document
from docx.oxml.ns import qn

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SK_DIR = os.path.abspath(os.path.join(BASE_DIR, '..'))
INPUT = os.path.join(SK_DIR, 'esg_admin', 'drafts', 'resume_v10.docx')
OUTPUT = os.path.join(BASE_DIR, 'drafts', 'resume_v1.docx')

CSR_INTRO = (
    "전기공학 학사와 사회복지학 석사(사회적기업 전공)를 거쳐, "
    "비영리 법인에서 3년간 사회공헌사업 기획·운영 실무를 담당했습니다.\n\n"

    "■ 사회공헌 프로그램 기획·운영 경험\n"
    "(사)함께만드는세상에서 KDB나눔재단의 '시니어브리지 사회공헌아카데미' "
    "사업을 기획 단계부터 운영·종료까지 전 주기에 걸쳐 담당했습니다. "
    "교육생 모집·접수 기획, 지역사회 협력기관 10개+ 발굴·선정·계약, "
    "봉사 프로그램 일정 조율, 교육장 환경 조성, 강사 섭외까지 "
    "프로그램 운영 전반을 주도했습니다.\n\n"

    "■ 기관 연계·후원 관리 및 기부금 정산\n"
    "협력기관 관리 프로세스를 표준화하여 정산 보고서 작성 시간을 30% 단축했습니다. "
    "재단 지원금 집행 내역 관리, 기부금 수입·지출 대사, 연말 결산 보고서 작성 등 "
    "기부금·보조금 관리 업무를 수행했습니다. "
    "인프라지원팀 배정 예산 약 3.6억 원의 집행·통제를 담당한 경험(집행률 88.5%)도 보유하고 있습니다.\n\n"

    "■ 사내 플랫폼·시스템 운영\n"
    "전사 ERP 시스템 도입 프로젝트를 실무 전담하여, "
    "요구사항 정의부터 회계·인사 프로세스 설계, 검수까지 수행했습니다. "
    "사내 IT 인프라 운영 관리 경험을 바탕으로 플랫폼 운영 업무에도 기여할 수 있습니다.\n\n"

    "■ 지원 동기\n"
    "비영리 현장에서 사회공헌 프로그램을 직접 기획·운영하며, "
    "이 분야에서 전문성을 키우겠다는 확신을 갖게 되었습니다. "
    "SK그룹의 사회적 가치 추구 경영철학에 깊이 공감하며, "
    "현장 경험과 사회복지사 1급 자격을 바탕으로 "
    "사회공헌 프로그램 기획·운영 업무에 기여하고 싶습니다."
)


def replace_intro(doc, new_text):
    """Table 2의 자기소개서 셀(1,0) 내용 교체 (서식 유지)"""
    t = doc.tables[2]
    cell = t.cell(1, 0)

    # 기존 텍스트의 서식 정보 보존
    first_para = cell.paragraphs[0]
    fmt = None
    if first_para.runs:
        run = first_para.runs[0]
        fmt = {
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
        }

    # 기존 내용 제거
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
        p.text = ''

    # 새 내용 입력 (첫 번째 paragraph에 전체 텍스트)
    first_para.text = ''
    run = first_para.add_run(new_text)
    if fmt:
        if fmt['font_name']:
            run.font.name = fmt['font_name']
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from lxml import etree
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), fmt['font_name'])
        if fmt['font_size']:
            run.font.size = fmt['font_size']
        if fmt['bold'] is not None:
            run.font.bold = fmt['bold']


def main():
    if not os.path.exists(INPUT):
        print(f"Error: {INPUT} not found")
        sys.exit(1)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

    doc = Document(INPUT)

    # 응시파트 변경: "ESG 사무" → "사회공헌 기획/운영"
    t0 = doc.tables[0]
    for row in t0.rows:
        for cell in row.cells:
            if 'ESG 사무' in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if 'ESG 사무' in run.text:
                            run.text = run.text.replace('ESG 사무', '사회공헌 기획/운영')
                            print("응시파트 변경: ESG 사무 → 사회공헌 기획/운영")

    replace_intro(doc, CSR_INTRO)
    doc.save(OUTPUT)
    print(f"저장: {OUTPUT}")
    print("자소서 교체 완료 (사회공헌 기획/운영 맞춤)")


if __name__ == '__main__':
    main()
